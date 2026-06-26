"""
Build the Isberian client report as a nicely-formatted .docx.

Output: isberian-report-2026-06-27.docx (in repo root).

Run with: python scripts/build-isberian-report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- styling tokens (mirror the site palette) ---
INK = RGBColor(0x1F, 0x1F, 0x1E)
MUTED = RGBColor(0x6E, 0x6E, 0x6C)
ACCENT = RGBColor(0x6B, 0x1F, 0x1A)  # oxblood
SUBTLE = RGBColor(0x88, 0x84, 0x7E)

SERIF = "Cormorant Garamond"
SANS = "Calibri"  # safe default present on most systems; Inter would be ideal but not guaranteed


def add_hr(doc):
    """Add a horizontal rule (a thin bottom border on an empty paragraph)."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9C9C5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)


def set_run(run, *, font=SANS, size=11, color=INK, bold=False, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    # Also set East Asian font so Word doesn't flip back to Calibri
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), font)
    r_fonts.set(qn("w:cs"), font)


def add_heading(doc, text, *, level=1):
    sizes = {1: 30, 2: 18, 3: 13}
    fonts = {1: SERIF, 2: SERIF, 3: SANS}
    colors = {1: INK, 2: INK, 3: ACCENT}
    spaces_before = {1: 18, 2: 16, 3: 10}
    spaces_after = {1: 10, 2: 6, 3: 4}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before[level])
    p.paragraph_format.space_after = Pt(spaces_after[level])
    run = p.add_run(text)
    set_run(
        run,
        font=fonts[level],
        size=sizes[level],
        color=colors[level],
        bold=(level == 3),
    )
    if level == 3:
        # Letter-spaced eyebrow style
        r_pr = run._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), "40")
        r_pr.append(spacing)


def add_para(doc, parts, *, size=11, after=8, italic=False, align=None, color=INK):
    """
    `parts` is either a string or a list of (text, {"bold": True/"italic": True/"em": True}) tuples
    for inline emphasis.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.35
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opts in parts:
        run = p.add_run(text)
        set_run(
            run,
            size=size,
            color=color,
            bold=bool(opts.get("bold")),
            italic=bool(opts.get("italic") or italic),
        )


def add_bullet(doc, parts, *, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    if isinstance(parts, str):
        parts = [(parts, {})]
    for text, opts in parts:
        run = p.add_run(text)
        set_run(
            run,
            size=size,
            color=INK,
            bold=bool(opts.get("bold")),
            italic=bool(opts.get("italic")),
        )


def main():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = SANS
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # --- Cover block ---
    add_para(
        doc,
        [("OSCAR ISBERIAN RUGS  ·  REPORT", {})],
        size=9,
        after=4,
        color=MUTED,
    )
    add_heading(doc, "Isberian on the modern web", level=1)
    add_para(
        doc,
        [("A report on the new site, in plain language", {"italic": True})],
        size=14,
        after=2,
        color=MUTED,
    )
    add_para(
        doc,
        [("Prepared for Oscar Isberian Rugs  ·  2026-06-27", {})],
        size=10,
        after=12,
        color=SUBTLE,
    )
    add_hr(doc)

    # --- The bottom line ---
    add_heading(doc, "The bottom line", level=2)
    add_para(
        doc,
        "For a hundred years the way to know Oscar Isberian Rugs has been to walk into the showroom. That doesn't change. What changes is how people find you before they walk in.",
    )
    add_para(
        doc,
        [
            ("A buyer today asks an AI assistant — Claude, ChatGPT, Perplexity, Google's AI Overviews — ", {}),
            ("“Where can I find an antique Heriz in Chicago?”", {"italic": True}),
            (" The current isberian.com loads its inventory through a browser script, which means assistants can't read it. To them, your inventory page reads ", {}),
            ("“Total Results: 0.”", {"italic": True}),
            (" A hundred years of catalog, invisible.", {}),
        ],
    )
    add_para(
        doc,
        "The new site fixes that without changing how you do business. Every rug is in the page the way a human reads it — so AI assistants and search engines can recommend you. Every visit converts to the same human outcome: a phone call, a showroom appointment, or a quote request. We never publish a price, we never invent stock, and we never tell anyone to clean a silk rug with dish soap.",
    )
    add_hr(doc)

    # --- What's actually new ---
    add_heading(doc, "What's actually new", level=2)

    features = [
        (
            "1.  You can be found by AI agents",
            [
                "The single biggest change. The new site is server-rendered — every catalog page returns the full content the moment it's requested, with no browser script needed. AI assistants read it, summarize it, and direct people to you.",
                "Each rug carries structured information in a format AI agents understand natively: title, origin, age, size, materials, weave, condition, color palette. They can answer questions like “Which Chicago dealer has antique Caucasian rugs under 6×9?” by reading your catalog directly.",
            ],
            "The next generation of customers will increasingly start their search by asking an assistant, not by typing into Google. If your site is invisible to those assistants, you're invisible to that audience. The new site is built so you show up when it matters.",
        ),
        (
            "2.  A concierge that knows the house voice",
            [
                "A small chat surface sits in the corner of every page. It's powered by Claude, and it speaks in your house voice — warm, precise, unhurried, never pushy. It knows:",
            ],
            "People land on your site at 11 PM. They have a question — “is this rug appropriate for a hallway with a dog?” In the old world, that question went unanswered and the visit didn't happen. The concierge answers in your voice and books the showroom visit.",
            [
                "The catalog (it can name only rugs that actually exist on your floor).",
                "Rug care, in restrained language (it will never recommend household cleaners on a silk or antique piece).",
                "The showrooms, the phones, and the booking path.",
                "When to stop and route a serious question to a human — “let me connect you with a specialist.”",
            ],
        ),
        (
            "3.  Photo triage for rug care and identification",
            [
                "Two related surfaces, both vision-based:",
            ],
            "These surfaces turn idle curiosity into a service inquiry without committing your time. The triage is always preliminary. The expertise stays in the showroom where it belongs.",
            [
                "/identify — a customer uploads photos of a piece they own. Claude returns a preliminary read of origin, age band, and type. Always preliminary. The report ends with a clear handoff to your in-showroom appraiser. We never appraise, value, or guarantee authenticity from photos.",
                "/services/triage — a customer photographs damage (wear, moth, fringe, stain) and adds a short note. Claude returns the issue and severity band, then routes them — drop-off, house call, or showroom inspection. Never DIY for valuable, antique, or silk pieces.",
            ],
        ),
        (
            "4.  Real rug pages, structured by hand",
            [
                "Every rug page renders a structured description block: a short lead paragraph, the technical details, the color palette as visible chips, the design features, what distinguishes the piece, and the provenance. AI drafts each block from the catalog data — an editor verifies before publish, and any unverified origin/age/knot-count claim stays visibly flagged.",
                "No empty superlatives. No “exquisite … masterpiece.” Specificity, restraint.",
            ],
            "Rug descriptions on most sites are interchangeable marketing prose. Yours are factual, specific, and machine-readable — which means search engines and AI agents understand them, surface them, and direct serious buyers to them.",
        ),
        (
            "5.  A care knowledge base, grounded in your practice",
            [
                "A library of short articles on care, materials, sizing, services, logistics, and the quote process. The concierge cites them when it answers care questions, so we never improvise. Antique and silk pieces always route to a specialist, never DIY.",
            ],
            "Customers searching for “how do I clean a Persian rug” land on real, restrained guidance — written in your voice — and book a service appointment rather than ruining the rug. Each article doubles as the substrate AI agents draw on when someone asks them the same question.",
        ),
        (
            "6.  The showroom path, on every page",
            [
                "A book-a-visit link, a Chicago phone, an Evanston phone — visible on every surface. The concierge mentions them too. There is no path on the site that doesn't end in a human option.",
            ],
            "The business model is built on the showroom visit. Every site decision points toward it.",
        ),
        (
            "7.  Crawlable previews on social and chat",
            [
                "When a customer shares a rug page on iMessage, Slack, Pinterest, LinkedIn — the share preview shows the rug itself, photographed against your brand mark, sized correctly. The old site shared the grey logo for every page. We share the rug.",
            ],
            "Quiet but compounding. Every shared link is a small ad. Showing the rug instead of the logo turns shares into clicks.",
        ),
    ]

    for entry in features:
        title = entry[0]
        paras = entry[1]
        why = entry[2]
        bullets = entry[3] if len(entry) > 3 else None

        add_heading(doc, title, level=3)
        for body in paras:
            add_para(doc, body)
        if bullets:
            for b in bullets:
                add_bullet(doc, b)
        add_para(
            doc,
            [("Why it matters.   ", {"bold": True}), (why, {})],
            after=14,
        )

    add_hr(doc)

    # --- What stays the same ---
    add_heading(doc, "What stays exactly the same", level=2)
    for line in [
        "The business model. Quoted only. No public prices. No anonymous checkout. No consumer accounts beyond email-based wishlists.",
        "The heritage tone. Master dealer, century of family practice. No emoji. No hype.",
        "The catalog truth. Only real rugs, with the real upstream record.",
        "The showrooms — both — with real hours and real phones.",
        "The five rules that gate every AI surface: no prices, no fabricated inventory, no valuations, no risky DIY, always a visible human exit.",
    ]:
        add_bullet(doc, line)

    add_hr(doc)

    # --- What we did NOT build ---
    add_heading(doc, "What we deliberately did not build", level=2)
    add_para(doc, "These are designed and waiting. Phase 2 or 3.", after=6)
    for line in [
        "Room visualizer — drop a rug into a photograph of your room. Needs the color-accurate photography pass.",
        "Visual search — “show me rugs that go with this fabric swatch.” Same photography prerequisite.",
        "Trade portal — designer logins, holds, memos. Needs operations process design.",
        "Custom rug studio — generative drafts for commissioned work.",
    ]:
        add_bullet(doc, line)
    add_para(
        doc,
        "We didn't pull them forward without sign-off. They're real options for next year.",
        after=10,
    )

    add_hr(doc)

    # --- The dependency ---
    add_heading(doc, "The single dependency in the way", level=2)
    add_para(
        doc,
        "The current isberian.com catalog lives in an external inventory system. The new site has been built against a small set of representative rugs (46 of them, real records, real photos), with a typed data layer ready to connect.",
    )
    add_para(
        doc,
        "For the new site to show the full catalog, that upstream system needs to expose a feed or be replicated into a dedicated catalog database. We've documented the data contract, written the connector, and stubbed both paths. The work to take it from 46 rugs to your full inventory is one focused engagement once the data path is decided.",
    )
    add_para(
        doc,
        [
            ("This is the only blocking dependency between the MVP and a launch you'd be comfortable with.", {"bold": True}),
        ],
    )

    add_hr(doc)

    # --- Why this matters ---
    add_heading(doc, "Why this matters, in one paragraph", level=2)
    add_para(
        doc,
        "The world that bought rugs from a phonebook has been replaced by a world that asks an assistant. The new isberian.com is the first version of the site that can answer that assistant honestly, in your voice, without compromising the showroom-first business that's worked for a century. The five rules that govern your floor now govern every line of code. Nothing about the way you sell rugs has changed. Everything about the way buyers can find you has.",
    )

    add_hr(doc)

    # --- What's live now ---
    add_heading(doc, "What's live now", level=2)
    add_para(
        doc,
        [
            ("You can see and use everything described above at ", {}),
            ("https://isberian-site-qbj6.vercel.app", {"bold": True}),
            (". Until DNS lands at the real domain, that's the address. Every change ships there within two minutes.", {}),
        ],
    )
    add_para(
        doc,
        "The catalog at the moment shows 46 rugs (curated from your existing photos). When the upstream feed lands, the same site will show the full inventory automatically — no rebuild, no rewrite.",
    )

    # --- Save ---
    out = os.path.join(os.path.dirname(__file__), "..", "isberian-report-2026-06-27.docx")
    out = os.path.normpath(out)
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
